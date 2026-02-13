import json
import os
import re
from datetime import datetime
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from tkinter.scrolledtext import ScrolledText
from collections import defaultdict


class SarifViewer(tk.Tk):

    def __init__(self):
        super().__init__()
        self.title("SARIF Viewer")
        self.geometry("1100x700")

        self.findings = []  # список сработок (dict)
        self.current_sarif_path: str | None = None
        # Файл для сохранения статусов сработок между запусками
        self._state_file = os.path.join(os.path.dirname(__file__), "sarif_viewer_state.json")
        self._state: dict[str, dict[str, str]] = {}
        self._load_state()

        # Тема оформления (светлая / тёмная)
        self.is_dark_theme: bool = False
        self._style = ttk.Style(self)
        try:
            self._default_ttk_theme = self._style.theme_use()
        except Exception:
            self._default_ttk_theme = None

        # Опциональная поддержка формирования DOCX‑отчёта
        self._docx = None
        try:
            import docx as _docx  # type: ignore
            self._docx = _docx
        except Exception:
            self._docx = None

        # Опциональная поддержка подсветки синтаксиса (Pygments)
        self._pygments_available = False
        self._pygments_lexers = {}
        try:
            from pygments.lexers import JavaLexer, ScalaLexer, JsonLexer, IniLexer, XmlLexer  # type: ignore

            self._pygments_available = True
            self._pygments_lexers = {
                "java": JavaLexer,
                "scala": ScalaLexer,
                "json": JsonLexer,
                "xml": XmlLexer,
                "conf": IniLexer,
                "ini": IniLexer,
                "cfg": IniLexer,
                "properties": IniLexer,
            }
        except Exception:
            self._pygments_available = False
            self._pygments_lexers = {}

        # Флаг видимости панели описания
        self.description_visible: bool = True

        # Опциональная поддержка Markdown в описании:
        # - markdown: конвертация Markdown -> HTML
        # - tkinterweb: HTML виджет для Tkinter
        self._md = None
        self._HtmlFrame = None
        try:
            import markdown as _markdown  # type: ignore
            self._md = _markdown
        except Exception:
            self._md = None
        try:
            from tkinterweb import HtmlFrame as _HtmlFrame  # type: ignore
            self._HtmlFrame = _HtmlFrame
        except Exception:
            self._HtmlFrame = None

        self._build_ui()

        # Обработчик закрытия окна — сохраняем состояние перед выходом
        try:
            self.protocol("WM_DELETE_WINDOW", self._on_close)
        except tk.TclError:
            pass

    # ---------------- UI ---------------- #
    def _build_ui(self):
        # Верхняя панель с кнопкой открытия SARIF
        top_frame = ttk.Frame(self)
        top_frame.pack(side=tk.TOP, fill=tk.X, padx=5, pady=5)

        open_btn = ttk.Button(top_frame, text="Открыть SARIF файл", command=self.open_sarif_file)
        open_btn.pack(side=tk.LEFT)

        load_json_btn = ttk.Button(top_frame, text="Загрузить JSON состояние", command=self.load_json_state)
        load_json_btn.pack(side=tk.LEFT, padx=(10, 0))

        compare_json_btn = ttk.Button(top_frame, text="Сравнить JSON файлы", command=self.compare_json_files)
        compare_json_btn.pack(side=tk.LEFT, padx=(10, 0))

        report_btn = ttk.Button(top_frame, text="Сформировать DOC отчёт", command=self.generate_doc_report)
        report_btn.pack(side=tk.LEFT, padx=(10, 0))

        # Переключатель темы (светлая/тёмная)
        self.theme_btn = ttk.Button(top_frame, text="Тёмная тема (experimental)", command=self.toggle_theme)
        self.theme_btn.pack(side=tk.LEFT, padx=(10, 0))

        self.current_file_label = ttk.Label(top_frame, text="Файл не выбран")
        self.current_file_label.pack(side=tk.LEFT, padx=10)

        # Основная панель с разделением (лево/право)
        main_pane = ttk.Panedwindow(self, orient=tk.HORIZONTAL)
        main_pane.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Левая панель – список сработок
        left_frame = ttk.Frame(main_pane)
        main_pane.add(left_frame, weight=1)

        # Заголовок + счётчики + кнопка сворачивания
        left_header_frame = ttk.Frame(left_frame)
        left_header_frame.pack(fill=tk.X)

        left_label = ttk.Label(left_header_frame, text="Сработки")
        left_label.pack(side=tk.LEFT, anchor="w")

        self.stats_label = ttk.Label(left_header_frame, text="", foreground="#555555")
        self.stats_label.pack(side=tk.LEFT, padx=(8, 0))

        collapse_btn = ttk.Button(left_header_frame, text="Свернуть все", command=self.collapse_all_rules)
        collapse_btn.pack(side=tk.RIGHT)

        # Группируем сработки по Rule id с раскрытием/сворачиванием
        self._tree_item_to_finding_index: dict[str, int] = {}
        self._rule_to_tree_item: dict[str, str] = {}

        style = self._style
        try:
            style.configure("Findings.Treeview", rowheight=30, font=("Segoe UI", 9))
        except tk.TclError:
            # на старых Tk параметры могут игнорироваться
            pass

        self.findings_tree = ttk.Treeview(
            left_frame,
            show="tree",
            selectmode="browse",
            style="Findings.Treeview",
        )
        self.findings_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.findings_tree.bind("<<TreeviewSelect>>", self.on_finding_selected)
        self.findings_tree.bind("<Double-1>", self.on_tree_double_click)

        scrollbar_left = ttk.Scrollbar(left_frame, orient=tk.VERTICAL, command=self.findings_tree.yview)
        scrollbar_left.pack(side=tk.RIGHT, fill=tk.Y)
        self.findings_tree.config(yscrollcommand=scrollbar_left.set)

        # Правая панель – код и описание
        right_frame = ttk.Frame(main_pane)
        main_pane.add(right_frame, weight=3)

        # Верхняя часть справа – код
        code_header_frame = ttk.Frame(right_frame)
        code_header_frame.pack(fill=tk.X)

        self.code_label = ttk.Label(code_header_frame, text="Фрагмент кода")
        self.code_label.pack(side=tk.LEFT, anchor="w")

        # Кнопка для скрытия/показа описания (справа сверху над окном кода)
        self.toggle_desc_btn = ttk.Button(
            code_header_frame,
            text="Скрыть описание",
            command=self.toggle_description_panel,
        )
        self.toggle_desc_btn.pack(side=tk.RIGHT)

        self.code_text = ScrolledText(right_frame, height=18, wrap=tk.NONE)
        self.code_text.pack(fill=tk.BOTH, expand=True)
        self.code_text.config(state=tk.DISABLED)
        # Тег для подсветки сработки в коде (бледно‑серый)
        try:
            self.code_text.tag_configure("finding_highlight", background="#e6e6e6")
        except tk.TclError:
            pass

        # Нижняя часть – описание и кнопки
        self.bottom_frame = ttk.Frame(right_frame)
        self.bottom_frame.pack(fill=tk.BOTH, expand=True, pady=(5, 0))

        # Кнопки подтверждения / отклонения
        buttons_frame = ttk.Frame(self.bottom_frame)
        buttons_frame.pack(fill=tk.X)

        reject_btn = ttk.Button(buttons_frame, text="Отклонить", command=self.reject_finding)
        reject_btn.pack(side=tk.LEFT, padx=5)

        confirm_btn = ttk.Button(buttons_frame, text="Подтвердить", command=self.confirm_finding)
        confirm_btn.pack(side=tk.LEFT, padx=5)

        undefined_btn = ttk.Button(buttons_frame, text="Не определено", command=self.undefined_finding)
        undefined_btn.pack(side=tk.LEFT, padx=5)

        reset_btn = ttk.Button(buttons_frame, text="Не обработано", command=self.reset_finding_status)
        reset_btn.pack(side=tk.LEFT, padx=5)

        # Выпадающий список для выбора критичности
        severity_label = ttk.Label(buttons_frame, text="Критичность:")
        severity_label.pack(side=tk.LEFT, padx=(20, 5))

        self.severity_combo = ttk.Combobox(
            buttons_frame,
            values=["Critical", "High", "Middle", "Low", "Error", "Info"],
            state="readonly",
            width=10
        )
        self.severity_combo.pack(side=tk.LEFT, padx=5)
        self.severity_combo.bind("<<ComboboxSelected>>", self.on_severity_changed)

        # Описание
        desc_label = ttk.Label(self.bottom_frame, text="Описание уязвимости")
        desc_label.pack(anchor="w", pady=(5, 0))

        # Описание: если доступны зависимости — рендерим Markdown как HTML,
        # иначе — отображаем как обычный текст (fallback).
        self.description_widget_kind = "text"
        self.description_text = None
        self.description_html = None
        if self._md and self._HtmlFrame:
            self.description_widget_kind = "html"
            self.description_html = self._HtmlFrame(self.bottom_frame, horizontal_scrollbar="auto")
            self.description_html.pack(fill=tk.BOTH, expand=True)
        else:
            self.description_text = ScrolledText(self.bottom_frame, height=10, wrap=tk.WORD)
            self.description_text.pack(fill=tk.BOTH, expand=True)
            self.description_text.config(state=tk.DISABLED)

        # Применяем светлую тему по умолчанию
        self._apply_light_theme()

    # ---------------- Работа с SARIF ---------------- #
    def open_sarif_file(self):
        # Сначала сохраняем состояние для уже открытого файла (если есть)
        self._save_current_file_state()

        filepath = filedialog.askopenfilename(
            title="Выберите SARIF файл",
            filetypes=[("SARIF / JSON файлы", "*.sarif *.json"), ("Все файлы", "*.*")]
        )
        if not filepath:
            return

        try:
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось прочитать файл:\n{e}")
            return

        full_path = os.path.abspath(filepath)
        self.current_sarif_path = full_path

        self.current_file_label.config(text=os.path.basename(filepath))
        self.load_findings_from_sarif(data, base_path=os.path.dirname(filepath))

    def load_findings_from_sarif(self, sarif_data, base_path: str | None = None):
        """
        Извлекает сработки из SARIF и заполняет список.
        """
        self.findings.clear()
        # очистка дерева
        for item in self.findings_tree.get_children(""):
            self.findings_tree.delete(item)
        self._tree_item_to_finding_index.clear()
        self._rule_to_tree_item.clear()

        runs = sarif_data.get("runs", [])
        for run in runs:
            results = run.get("results", [])
            for res in results:
                finding = self._parse_finding(res, base_path)
                self.findings.append(finding)

        # группировка по ruleId
        grouped: dict[str, list[int]] = defaultdict(list)
        for idx, f in enumerate(self.findings):
            grouped[f.get("ruleId") or "N/A"].append(idx)

        # порядок: по ruleId, внутри — как в исходном файле
        for rule_id in sorted(grouped.keys(), key=lambda s: str(s).lower()):
            count = len(grouped[rule_id])
            rule_item = self.findings_tree.insert("", tk.END, text=f"{rule_id} ({count})", open=False)
            self._rule_to_tree_item[rule_id] = rule_item
            for idx in grouped[rule_id]:
                f = self.findings[idx]
                title = f"{idx + 1}. [{f['level'].upper()}] {f['message']}"
                child = self.findings_tree.insert(rule_item, tk.END, text=title, open=False)
                self._tree_item_to_finding_index[child] = idx
                self.findings[idx]["tree_item"] = child

        # Восстанавливаем статусы из сохранённого состояния (если есть)
        self._apply_state_to_findings()

        # Обновляем все Rule id (счетчики и цвета)
        for rule_id in grouped.keys():
            self._update_rule_item(rule_id)

        self._recalc_stats()

    def _parse_finding(self, result: dict, base_path: str | None):
        """
        Преобразует один result из SARIF в удобный для GUI формат.
        """
        rule_id = result.get("ruleId", "N/A")
        message = (result.get("message") or {}).get("text", "")
        level = result.get("level", "warning")

        locations = result.get("locations") or []
        location = locations[0] if locations else {}
        physical = location.get("physicalLocation") or {}
        artifact = physical.get("artifactLocation") or {}
        region = physical.get("region") or {}

        uri = artifact.get("uri", "")
        start_line = region.get("startLine")
        end_line = region.get("endLine")
        snippet_text = (region.get("snippet") or {}).get("text", "")

        # Попытаемся загрузить полный файл, если он доступен
        full_code = None
        if uri:
            # URI в SARIF обычно относительный путь к проекту
            candidate_paths = []
            if base_path:
                candidate_paths.append(os.path.join(base_path, uri))
            candidate_paths.append(uri)

            for p in candidate_paths:
                if os.path.isfile(p):
                    try:
                        with open(p, "r", encoding="utf-8", errors="replace") as f:
                            full_code = f.read()
                        break
                    except Exception:
                        full_code = None

        properties = result.get("properties") or {}
        llm = properties.get("llmAnalysis") or {}

        description_parts = []

        # Основное объяснение (русское reasoning, если есть)
        if properties.get("reasoning"):
            description_parts.append("Анализ:\n" + properties["reasoning"])

        # Анализ из llmAnalysis (английский)
        if llm.get("reasoning"):
            description_parts.append("\nLLM‑анализ:\n" + llm["reasoning"])

        # Эксплуатационная цепочка (русская)
        exploitation_chain = properties.get("exploitationChain") or []
        if exploitation_chain:
            description_parts.append("\nЦепочка эксплуатации:")
            for step in exploitation_chain:
                num = step.get("step")
                desc = step.get("description", "")
                description_parts.append(f"  {num}. {desc}")

        # Цепочка из llmAnalysis
        vuln_chain = llm.get("vulnerabilityChain") or []
        if vuln_chain:
            description_parts.append("\nVulnerability chain (LLM):")
            for step in vuln_chain:
                num = step.get("step")
                desc = step.get("description", "")
                description_parts.append(f"  {num}. {desc}")

        # Рекомендации
        suggested_fix = properties.get("suggestedFix") or llm.get("suggestedFix")
        if suggested_fix:
            description_parts.append("\nРекомендуемое исправление:\n" + str(suggested_fix))

        # CVSS и статус
        cvss = properties.get("cvss31")
        if cvss:
            description_parts.append(f"\nCVSS: {cvss}")

        confirmation = properties.get("confirmationStatus")
        if confirmation:
            description_parts.append(f"Статус подтверждения: {confirmation}")

        description = "\n".join(description_parts).strip()

        # Извлекаем reasoning отдельно для использования в отчете
        reasoning = properties.get("reasoning") or ""
        # reasoning может быть объектом с полем "text"
        if isinstance(reasoning, dict):
            reasoning = reasoning.get("text", "") or ""
        
        # Извлекаем suggestedFix отдельно для использования в отчете
        # Проверяем несколько возможных мест в SARIF структуре
        suggested_fix_for_report = (
            properties.get("suggestedFix") or 
            llm.get("suggestedFix") or 
            ""
        )
        
        # Также проверяем result.get("fixes") - это массив объектов fix
        fixes = result.get("fixes")
        if not suggested_fix_for_report and fixes:
            # Если fixes - это массив, берем первый элемент
            if isinstance(fixes, list) and len(fixes) > 0:
                first_fix = fixes[0]
                if isinstance(first_fix, dict):
                    # В fix может быть description или другие поля
                    suggested_fix_for_report = (
                        first_fix.get("description", {}).get("text", "") or
                        first_fix.get("description", "") or
                        ""
                    )
        
        # Обрабатываем разные форматы suggestedFix
        if isinstance(suggested_fix_for_report, dict):
            # Если это объект, пытаемся извлечь text или description
            suggested_fix_for_report = (
                suggested_fix_for_report.get("text", "") or 
                suggested_fix_for_report.get("description", "") or 
                ""
            )
        elif isinstance(suggested_fix_for_report, list):
            # Если это список, объединяем все элементы
            if len(suggested_fix_for_report) > 0:
                # Если первый элемент - объект, извлекаем text
                if isinstance(suggested_fix_for_report[0], dict):
                    texts = []
                    for item in suggested_fix_for_report:
                        text = item.get("text") or item.get("description") or str(item)
                        if text:
                            texts.append(str(text))
                    suggested_fix_for_report = "\n".join(texts)
                else:
                    suggested_fix_for_report = "\n".join(str(item) for item in suggested_fix_for_report if item)
            else:
                suggested_fix_for_report = ""
        
        # Преобразуем в строку, если еще не строка
        if not isinstance(suggested_fix_for_report, str):
            suggested_fix_for_report = str(suggested_fix_for_report) if suggested_fix_for_report else ""

        # Определяем критичность по умолчанию из level
        severity_mapping = {
            "error": "Error",
            "warning": "Middle",
            "note": "Info"
        }
        default_severity = severity_mapping.get(level.lower(), "Middle")

        return {
            "ruleId": rule_id,
            "message": message,
            "level": level,
            "uri": uri,
            "startLine": start_line,
            "endLine": end_line,
            "snippet": snippet_text,
            "full_code": full_code,
            "description": description,
            "reasoning": reasoning,  # Отдельное поле для reasoning
            "suggestedFix": suggested_fix_for_report,  # Отдельное поле для suggestedFix
            "status": None,  # None / 'confirmed' / 'rejected' / 'undefined'
            "severity": default_severity,  # 'crit' / 'high' / 'mid' / 'low' / 'err' / 'info'
            "tree_item": None,
        }

    # ---------------- Обработчики GUI ---------------- #
    def on_finding_selected(self, event=None):
        selection = self.findings_tree.selection()
        if not selection:
            return
        item = selection[0]

        # Если выбрали Rule id — только раскрываем/сворачиваем
        if item not in self._tree_item_to_finding_index:
            self._toggle_rule_item(item)
            return

        idx = self._tree_item_to_finding_index[item]
        finding = self.findings[idx]

        # Обновляем код
        self._show_code_for_finding(finding)

        # Обновляем описание
        self._set_description(finding.get("description") or "Описание отсутствует.")

        # Обновляем выпадающий список критичности для выбранной сработки
        severity = finding.get("severity")
        if severity:
            try:
                self.severity_combo.set(severity)
            except tk.TclError:
                pass
        else:
            # Если критичность не установлена, пытаемся определить из level
            level = finding.get("level", "").lower()
            severity_mapping = {
                "error": "Error",
                "warning": "Middle",
                "note": "Info"
            }
            default_severity = severity_mapping.get(level, "Middle")
            # Устанавливаем значение по умолчанию в finding, если его еще нет
            if not finding.get("severity"):
                finding["severity"] = default_severity
            try:
                self.severity_combo.set(default_severity)
            except tk.TclError:
                pass

    def on_tree_double_click(self, event=None):
        # Даблклик по Rule id — тоже раскрывает/сворачивает
        item = self.findings_tree.focus()
        if not item:
            return
        if item not in self._tree_item_to_finding_index:
            self._toggle_rule_item(item)

    def on_severity_changed(self, event=None):
        """Обработчик изменения критичности в выпадающем списке"""
        idx = self._get_selected_index()
        if idx is None:
            return
        severity = self.severity_combo.get()
        if severity:
            self.findings[idx]["severity"] = severity
            # Сохраняем состояние сразу после изменения критичности
            self._save_current_file_state()
            # Обновляем цвет элемента в дереве (если нужно)
            self._update_listbox_item_color(idx)

    def _toggle_rule_item(self, item: str):
        try:
            is_open = bool(self.findings_tree.item(item, "open"))
            self.findings_tree.item(item, open=not is_open)
        except tk.TclError:
            pass

    def _show_code_for_finding(self, finding: dict):
        """
        Показывает код. Если доступен полный файл — показывает ВЕСЬ файл
        с подсветкой диапазона строк сработки (startLine–endLine).
        """
        # Обновляем заголовок с путём к файлу
        if finding.get("uri"):
            try:
                self.code_label.config(text=f"Фрагмент кода — {finding['uri']}")
            except tk.TclError:
                pass
        else:
            try:
                self.code_label.config(text="Фрагмент кода")
            except tk.TclError:
                pass

        header_parts = []
        if finding["uri"]:
            header_parts.append(f"Файл: {finding['uri']}")
        if finding["startLine"]:
            header_parts.append(f"Строки: {finding['startLine']}–{finding['endLine']}")

        header = " | ".join(header_parts)
        if finding.get("full_code"):
            self._set_code_with_highlight(
                widget=self.code_text,
                header=header,
                full_code=finding["full_code"],
                start_line=finding.get("startLine"),
                end_line=finding.get("endLine"),
                uri=finding.get("uri"),
            )
            return

        # Fallback: если полный код недоступен — показываем snippet (без подсветки по строкам)
        code_to_show = finding.get("snippet") or "Фрагмент кода недоступен."
        if header:
            code_to_show = header + "\n\n" + code_to_show
        self._set_text(self.code_text, code_to_show)

    def _set_text(self, widget: ScrolledText, text: str):
        if widget is None:
            return
        widget.config(state=tk.NORMAL)
        widget.delete("1.0", tk.END)
        widget.insert(tk.END, text)
        widget.config(state=tk.DISABLED)

    def _escape_html_tags_in_text(self, text: str) -> str:
        """
        Экранирует HTML-подобные теги в тексте перед обработкой markdown,
        но НЕ трогает теги внутри блоков кода (```...```), так как markdown
        сам правильно их обработает.
        Например, <root level="DEBUG"> будет экранирован, чтобы отображаться как текст,
        но внутри блоков кода теги останутся как есть.
        """
        # Разделяем текст на части: блоки кода и обычный текст
        # Паттерн для блоков кода markdown: ```язык\n...\n```
        code_block_pattern = r'```[\s\S]*?```'
        
        parts = []
        last_end = 0
        
        # Находим все блоки кода
        for match in re.finditer(code_block_pattern, text):
            # Текст до блока кода - экранируем HTML-теги
            before_code = text[last_end:match.start()]
            escaped_before = self._escape_html_tags_in_plain_text(before_code)
            parts.append(escaped_before)
            
            # Блок кода - оставляем как есть (markdown сам обработает)
            code_block = match.group(0)
            parts.append(code_block)
            
            last_end = match.end()
        
        # Остаток текста после последнего блока кода - экранируем HTML-теги
        remaining = text[last_end:]
        escaped_remaining = self._escape_html_tags_in_plain_text(remaining)
        parts.append(escaped_remaining)
        
        return ''.join(parts)
    
    def _escape_html_tags_in_plain_text(self, text: str) -> str:
        """
        Экранирует HTML-подобные теги в обычном тексте (не в блоках кода).
        """
        def escape_tag(match):
            full_tag = match.group(0)
            # Экранируем весь тег
            return full_tag.replace('<', '&lt;').replace('>', '&gt;')
        
        # Регулярное выражение для поиска HTML-тегов
        # Ищем < за которым следует имя тега (буквы/цифры) и опционально атрибуты, затем >
        html_tag_pattern = r'<[a-zA-Z][a-zA-Z0-9]*(?:\s+[^>]*)?>'
        return re.sub(html_tag_pattern, escape_tag, text)

    def _set_description(self, markdown_text: str):
        """
        Пишем описание сработки.
        - Если есть markdown+tkinterweb — рендерим Markdown.
        - Иначе отображаем исходный текст как есть.
        """
        if self.description_widget_kind == "html" and self.description_html is not None:
            # Экранируем HTML-теги в исходном тексте перед обработкой markdown,
            # чтобы они не интерпретировались как HTML-разметка
            # (например, <root level="DEBUG"> должен отображаться как текст)
            # Используем простое экранирование всех < и >, так как markdown
            # может удалить невалидные HTML-теги
            escaped_text = self._escape_html_tags_in_text(markdown_text)
            try:
                html = self._md.markdown(  # type: ignore[union-attr]
                    escaped_text,
                    extensions=["fenced_code", "tables", "toc"],
                    output_format="html5",
                )
            except Exception:
                html = "<pre>" + self._escape_html(markdown_text) + "</pre>"
            # Добавим минимальные стили, чтобы читалось аккуратно
            # CSS в зависимости от текущей темы
            if getattr(self, "is_dark_theme", False):
                css = (
                    "body{font-family:Segoe UI,Arial,sans-serif;font-size:10pt;"
                    "background:#1e1e1e;color:#d4d4d4;}"
                    "pre,code{font-family:Consolas,Menlo,monospace;font-size:9pt;}"
                    "pre{background:#252526;padding:8px;border-radius:6px;}"
                    "code{background:#333333;padding:1px 3px;border-radius:4px;}"
                    "table{border-collapse:collapse;}"
                    "td,th{border:1px solid #444;padding:6px;}"
                )
            else:
                css = (
                    "body{font-family:Segoe UI,Arial,sans-serif;font-size:10pt;}"
                    "pre,code{font-family:Consolas,Menlo,monospace;font-size:9pt;}"
                    "pre{background:#f6f6f6;padding:8px;border-radius:6px;}"
                    "code{background:#f0f0f0;padding:1px 3px;border-radius:4px;}"
                    "table{border-collapse:collapse;}td,th{border:1px solid #ddd;padding:6px;}"
                )

            html_doc = (
                "<html><head><meta charset='utf-8'>"
                "<style>"
                + css +
                "</style></head><body>"
                + html
                + "</body></html>"
            )
            try:
                self.description_html.load_html(html_doc)
            except Exception:
                # На случай, если HtmlFrame не смог — fallback в текст
                if self.description_text is not None:
                    self._set_text(self.description_text, markdown_text)
            return

        # text fallback
        if self.description_text is not None:
            self._set_text(self.description_text, markdown_text)

    def _escape_html(self, s: str) -> str:
        return (
            s.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&#39;")
        )

    def _set_code_with_highlight(
        self,
        widget: ScrolledText,
        header: str,
        full_code: str,
        start_line: int | None,
        end_line: int | None,
        uri: str | None,
    ):
        """
        Показывает ВЕСЬ файл с нумерацией строк и подсветкой диапазона строк.
        Подсветка делается тегом 'finding_highlight' с бледно‑серым фоном.
        """
        lines = full_code.splitlines()
        max_no = max(1, len(lines))
        width = len(str(max_no))

        # Нормализуем диапазон подсветки
        hl_start = None
        hl_end = None
        if isinstance(start_line, int) and start_line >= 1:
            hl_start = start_line
        if isinstance(end_line, int) and end_line >= 1:
            hl_end = end_line
        if hl_start is not None and hl_end is None:
            hl_end = hl_start
        if hl_start is not None and hl_end is not None and hl_end < hl_start:
            hl_start, hl_end = hl_end, hl_start

        widget.config(state=tk.NORMAL)
        widget.delete("1.0", tk.END)
        try:
            widget.tag_remove("finding_highlight", "1.0", tk.END)
        except tk.TclError:
            pass

        # Собираем текст целиком
        text_parts: list[str] = []
        header_lines_count = 0
        if header:
            text_parts.append(header)
            text_parts.append("")  # пустая строка
            header_lines_count = 2

        numbered_lines = [f"{i:{width}d}: {line}" for i, line in enumerate(lines, start=1)]
        text_parts.extend(numbered_lines)

        full_text = "\n".join(text_parts) + ("\n" if lines else "")
        widget.insert("1.0", full_text)

        # Подсветка нужных строк: считаем, что первая строка кода
        # начинается после header_lines_count строк.
        if hl_start is not None and hl_end is not None:
            code_start_widget_line = header_lines_count + 1
            for i in range(hl_start, hl_end + 1):
                widget_line = code_start_widget_line + (i - 1)
                start_index = f"{widget_line}.0"
                end_index = f"{widget_line}.end"
                try:
                    widget.tag_add("finding_highlight", start_index, end_index)
                except tk.TclError:
                    pass

            # Прокручиваем так, чтобы первая подсвеченная строка была видна
            first_widget_line = code_start_widget_line + (hl_start - 1)
            try:
                widget.see(f"{first_widget_line}.0")
            except tk.TclError:
                pass

        # Подсветка синтаксиса (если доступна)
        self._apply_syntax_highlighting(
            widget=widget,
            full_code=full_code,
            header_lines_count=header_lines_count,
            line_number_width=width,
            uri=uri,
        )

        widget.config(state=tk.DISABLED)

    def _detect_language_from_uri(self, uri: str | None) -> str | None:
        if not uri:
            return None
        uri_lower = uri.lower()
        if uri_lower.endswith(".java"):
            return "java"
        if uri_lower.endswith(".scala") or uri_lower.endswith(".sc") or uri_lower.endswith(".sbt"):
            return "scala"
        if uri_lower.endswith(".json"):
            return "json"
        if uri_lower.endswith(".xml"):
            return "xml"
        if uri_lower.endswith((".conf", ".ini", ".cfg", ".properties")):
            return "conf"
        # Для других типов (csv и пр.) оставим без подсветки
        return None

    def _apply_syntax_highlighting(
        self,
        widget: ScrolledText,
        full_code: str,
        header_lines_count: int,
        line_number_width: int,
        uri: str | None,
    ):
        """
        Подсвечивает синтаксис для Java/Scala/JSON/конфигов, если установлен pygments.
        """
        if not self._pygments_available or not self._pygments_lexers:
            return

        lang = self._detect_language_from_uri(uri)
        if not lang:
            return

        LexerCls = self._pygments_lexers.get(lang)
        if not LexerCls:
            return

        try:
            from pygments import lex  # type: ignore
            from pygments.token import Token  # type: ignore
        except Exception:
            return

        try:
            lexer = LexerCls()
        except Exception:
            return

        # Настраиваем теги стилей (однократно)
        try:
            widget.tag_configure("syn_keyword", foreground="#0000cc")
            widget.tag_configure("syn_type", foreground="#267f99")
            widget.tag_configure("syn_string", foreground="#a31515")
            widget.tag_configure("syn_comment", foreground="#6a9955")
            widget.tag_configure("syn_number", foreground="#098658")
        except tk.TclError:
            pass

        code_start_widget_line = header_lines_count + 1
        prefix_len = line_number_width + 2  # "NNN: "

        lines = full_code.splitlines()
        for i, line in enumerate(lines, start=1):
            widget_line = code_start_widget_line + (i - 1)
            col = 0
            try:
                tokens = list(lex(line, lexer))
            except Exception:
                continue

            for ttype, value in tokens:
                length = len(value)
                if length == 0:
                    continue

                # Определяем тег по типу токена
                tag = None
                if ttype in Token.Keyword or str(ttype).startswith("Token.Keyword"):
                    tag = "syn_keyword"
                elif ttype in Token.Name.Class or str(ttype).startswith("Token.Name.Class"):
                    tag = "syn_type"
                elif ttype in Token.Name.Function or str(ttype).startswith("Token.Name.Function"):
                    tag = "syn_type"
                elif ttype in Token.String or str(ttype).startswith("Token.String"):
                    tag = "syn_string"
                elif ttype in Token.Comment or str(ttype).startswith("Token.Comment"):
                    tag = "syn_comment"
                elif ttype in Token.Literal.Number or str(ttype).startswith("Token.Literal.Number"):
                    tag = "syn_number"

                if tag:
                    start_index = f"{widget_line}.{prefix_len + col}"
                    end_index = f"{widget_line}.{prefix_len + col + length}"
                    try:
                        widget.tag_add(tag, start_index, end_index)
                    except tk.TclError:
                        pass

                col += length

    def _update_listbox_item_color(self, idx: int):
        finding = self.findings[idx]
        status = finding["status"]

        # Сброс к стандартному виду
        fg = "black"
        bg = "white"

        if status == "confirmed":
            bg = "#ffcccc"  # светло‑красный
            fg = "black"
        elif status == "rejected":
            bg = "#ccffcc"  # светло‑зелёный
            fg = "black"
        elif status == "undefined":
            bg = "#ffffcc"  # светло‑жёлтый
            fg = "black"

        # Обновляем цвет строки в дереве (только для конкретной сработки)
        item = finding.get("tree_item")
        if item:
            tag = f"status_{status or 'none'}"
            try:
                self.findings_tree.tag_configure(tag, background=bg, foreground=fg)
                self.findings_tree.item(item, tags=(tag,))
            except tk.TclError:
                pass
        
        # Обновляем Rule id (счетчик и цвет, если все обработано)
        rule_id = finding.get("ruleId")
        if rule_id:
            self._update_rule_item(rule_id)
        
        self._recalc_stats()

    def _update_rule_item(self, rule_id: str):
        """
        Обновляет Rule id в дереве: счетчик сработок и цвет фона
        (светло-серый, если все сработки обработаны; светло-жёлтый, если есть "не определено").
        """
        rule_item = self._rule_to_tree_item.get(rule_id)
        if not rule_item:
            return

        # Находим все сработки для этого rule_id
        rule_findings = [f for f in self.findings if f.get("ruleId") == rule_id]
        total_count = len(rule_findings)
        
        if total_count == 0:
            return

        # Проверяем наличие сработок со статусом "undefined"
        has_undefined = any(f.get("status") == "undefined" for f in rule_findings)
        
        # Подсчитываем обработанные (confirmed или rejected)
        processed_count = sum(1 for f in rule_findings if f.get("status") in ("confirmed", "rejected"))
        all_processed = processed_count == total_count

        # Обновляем текст с актуальным счетчиком
        new_text = f"{rule_id} ({total_count})"
        try:
            self.findings_tree.item(rule_item, text=new_text)
        except tk.TclError:
            pass

        # Настраиваем цвет фона
        try:
            if has_undefined:
                # Светло-жёлтый фон для правил с "не определено"
                tag_name = "rule_has_undefined"
                bg_color = "#ffffcc" if not getattr(self, "is_dark_theme", False) else "#5a5a3a"
                self.findings_tree.tag_configure(tag_name, background=bg_color)
                self.findings_tree.item(rule_item, tags=(tag_name,))
            elif all_processed:
                # Светло-серый фон для полностью обработанных правил
                tag_name = "rule_all_processed"
                bg_color = "#e0e0e0" if not getattr(self, "is_dark_theme", False) else "#3a3a3a"
                self.findings_tree.tag_configure(tag_name, background=bg_color)
                self.findings_tree.item(rule_item, tags=(tag_name,))
            else:
                # Удаляем тег, чтобы использовался фон по умолчанию дерева
                self.findings_tree.item(rule_item, tags=())
        except tk.TclError:
            pass

    def _get_selected_index(self) -> int | None:
        selection = self.findings_tree.selection()
        if not selection:
            messagebox.showinfo("Нет выбора", "Сначала выберите сработку в списке слева.")
            return None
        item = selection[0]
        if item not in self._tree_item_to_finding_index:
            messagebox.showinfo("Нет выбора", "Выберите конкретную сработку (строку под Rule id).")
            return None
        return self._tree_item_to_finding_index[item]

    # --------- Работа с сохранением состояния --------- #
    def _make_finding_key(self, f: dict) -> str:
        """
        Делает стабильный ключ для идентификации сработки внутри SARIF‑файла.
        """
        return "|".join(
            [
                str(f.get("ruleId") or ""),
                str(f.get("uri") or ""),
                str(f.get("startLine") or ""),
                str(f.get("endLine") or ""),
                str(f.get("message") or ""),
            ]
        )

    def _load_state(self):
        try:
            if os.path.isfile(self._state_file):
                with open(self._state_file, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if isinstance(data, dict):
                    self._state = data  # type: ignore[assignment]
        except Exception:
            self._state = {}

    def _save_state(self):
        try:
            with open(self._state_file, "w", encoding="utf-8") as f:
                json.dump(self._state, f, ensure_ascii=False, indent=2)
        except Exception:
            # Не мешаем работе GUI, если не удалось сохранить
            pass

    def _save_current_file_state(self):
        """
        Сохраняет статусы и критичность для текущего SARIF‑файла в общий state.
        """
        if not self.current_sarif_path or not self.findings:
            return
        file_key = self.current_sarif_path
        file_state = self._state.setdefault(file_key, {})

        for f in self.findings:
            key = self._make_finding_key(f)
            status = f.get("status")
            severity = f.get("severity")
            
            state_entry = {}
            if status in ("confirmed", "rejected", "undefined"):
                state_entry["status"] = status
            if severity:
                state_entry["severity"] = severity
            
            if state_entry:
                file_state[key] = state_entry
            else:
                file_state.pop(key, None)

        self._save_state()

    def _apply_state_to_findings(self):
        """
        Применяет сохранённые статусы и критичность к сработкам после загрузки SARIF‑файла.
        """
        if not self.current_sarif_path:
            return
        file_state = self._state.get(self.current_sarif_path) or {}
        if not file_state:
            return

        for idx, f in enumerate(self.findings):
            key = self._make_finding_key(f)
            state_entry = file_state.get(key)
            if isinstance(state_entry, dict):
                # Новый формат с severity
                if state_entry.get("status") in ("confirmed", "rejected", "undefined"):
                    f["status"] = state_entry["status"]
                if state_entry.get("severity"):
                    f["severity"] = state_entry["severity"]
            elif isinstance(state_entry, str):
                # Старый формат (только status)
                if state_entry in ("confirmed", "rejected", "undefined"):
                    f["status"] = state_entry
            self._update_listbox_item_color(idx)

    def _on_close(self):
        """
        Обработчик закрытия главного окна: сохраняем состояние и закрываем приложение.
        """
        self._save_current_file_state()
        self.destroy()

    def load_json_state(self):
        """
        Загружает JSON файл состояния для текущего SARIF файла.
        Позволяет коллеге открыть JSON файл пользователя в программе.
        """
        if not self.current_sarif_path:
            messagebox.showwarning("Нет SARIF файла", "Сначала откройте SARIF файл.")
            return

        filepath = filedialog.askopenfilename(
            title="Выберите JSON файл состояния",
            filetypes=[("JSON файлы", "*.json"), ("Все файлы", "*.*")]
        )
        if not filepath:
            return

        try:
            with open(filepath, "r", encoding="utf-8") as f:
                loaded_state = json.load(f)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось прочитать JSON файл:\n{e}")
            return

        if not isinstance(loaded_state, dict):
            messagebox.showerror("Ошибка", "JSON файл должен содержать словарь (dict).")
            return

        # Ищем состояние для текущего SARIF файла в загруженном JSON
        # Проверяем, есть ли ключ с путем к текущему SARIF файлу
        file_key = self.current_sarif_path
        file_state = loaded_state.get(file_key)

        # Если точного совпадения нет, пытаемся найти по имени файла
        if not file_state:
            sarif_basename = os.path.basename(file_key)
            for key, value in loaded_state.items():
                if os.path.basename(key) == sarif_basename:
                    file_key = key
                    file_state = value
                    break

        if not file_state:
            messagebox.showwarning(
                "Не найдено",
                f"В выбранном JSON файле не найдено состояние для текущего SARIF файла:\n{os.path.basename(self.current_sarif_path)}\n\n"
                "Убедитесь, что JSON файл был создан для этого SARIF файла."
            )
            return

        # Применяем загруженное состояние
        if not isinstance(file_state, dict):
            messagebox.showerror("Ошибка", "Состояние для SARIF файла должно быть словарем (dict).")
            return

        # Обновляем состояние в памяти
        self._state[file_key] = file_state.copy()
        
        # Применяем состояние к текущим сработкам
        for idx, f in enumerate(self.findings):
            key = self._make_finding_key(f)
            state_entry = file_state.get(key)
            if isinstance(state_entry, dict):
                if state_entry.get("status") in ("confirmed", "rejected", "undefined"):
                    f["status"] = state_entry["status"]
                if state_entry.get("severity"):
                    f["severity"] = state_entry["severity"]
            elif isinstance(state_entry, str):
                if state_entry in ("confirmed", "rejected", "undefined"):
                    f["status"] = state_entry
            self._update_listbox_item_color(idx)

        # Обновляем все Rule id
        for rule_id in set(f.get("ruleId") for f in self.findings if f.get("ruleId")):
            self._update_rule_item(rule_id)

        self._recalc_stats()
        messagebox.showinfo("Готово", "JSON состояние успешно загружено и применено.")

    def compare_json_files(self):
        """
        Сравнивает два JSON файла состояния одного SARIF файла и открывает окно сравнения.
        В окне сравнения показываются только обработанные сработки (те, что отличаются).
        """
        if not self.current_sarif_path:
            messagebox.showwarning("Нет SARIF файла", "Сначала откройте SARIF файл.")
            return

        # Запрашиваем первый JSON файл
        filepath1 = filedialog.askopenfilename(
            title="Выберите первый JSON файл состояния",
            filetypes=[("JSON файлы", "*.json"), ("Все файлы", "*.*")]
        )
        if not filepath1:
            return

        # Запрашиваем второй JSON файл
        filepath2 = filedialog.askopenfilename(
            title="Выберите второй JSON файл состояния",
            filetypes=[("JSON файлы", "*.json"), ("Все файлы", "*.*")]
        )
        if not filepath2:
            return

        if filepath1 == filepath2:
            messagebox.showwarning("Ошибка", "Выберите два разных JSON файла.")
            return

        try:
            with open(filepath1, "r", encoding="utf-8") as f:
                state1 = json.load(f)
            with open(filepath2, "r", encoding="utf-8") as f:
                state2 = json.load(f)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось прочитать JSON файлы:\n{e}")
            return

        if not isinstance(state1, dict) or not isinstance(state2, dict):
            messagebox.showerror("Ошибка", "JSON файлы должны содержать словари (dict).")
            return

        # Ищем состояние для текущего SARIF файла в обоих JSON
        file_key = self.current_sarif_path
        file_state1 = state1.get(file_key)
        file_state2 = state2.get(file_key)

        # Если точного совпадения нет, пытаемся найти по имени файла
        if not file_state1:
            sarif_basename = os.path.basename(file_key)
            for key, value in state1.items():
                if os.path.basename(key) == sarif_basename:
                    file_key = key
                    file_state1 = value
                    break

        if not file_state2:
            sarif_basename = os.path.basename(file_key)
            for key, value in state2.items():
                if os.path.basename(key) == sarif_basename:
                    file_state2 = value
                    break

        if not file_state1 or not file_state2:
            messagebox.showwarning(
                "Не найдено",
                f"В одном из JSON файлов не найдено состояние для текущего SARIF файла:\n{os.path.basename(self.current_sarif_path)}"
            )
            return

        # Открываем окно сравнения
        comparison_window = ComparisonWindow(
            self,
            self.findings,
            file_state1,
            file_state2,
            os.path.basename(filepath1),
            os.path.basename(filepath2),
            self.current_sarif_path
        )

    def _recalc_stats(self):
        """
        Пересчитывает количество подтверждённых / отклонённых / неопределённых / необработанных сработок
        и обновляет счётчик над списком.
        """
        confirmed = 0
        rejected = 0
        undefined = 0
        pending = 0
        for f in self.findings:
            if f["status"] == "confirmed":
                confirmed += 1
            elif f["status"] == "rejected":
                rejected += 1
            elif f["status"] == "undefined":
                undefined += 1
            else:
                pending += 1

        text = f"Подтв.: {confirmed}  Откл.: {rejected}  Не опр.: {undefined}  Не обработано: {pending}"
        try:
            self.stats_label.config(text=text)
        except tk.TclError:
            pass

    def collapse_all_rules(self):
        """
        Сворачивает все раскрытые Rule id в дереве.
        """
        for item in self.findings_tree.get_children(""):
            try:
                self.findings_tree.item(item, open=False)
            except tk.TclError:
                pass

    def reject_finding(self):
        idx = self._get_selected_index()
        if idx is None:
            return
        self._set_finding_status(idx, "rejected")

    def confirm_finding(self):
        idx = self._get_selected_index()
        if idx is None:
            return
        self._set_finding_status(idx, "confirmed")

    def undefined_finding(self):
        idx = self._get_selected_index()
        if idx is None:
            return
        self._set_finding_status(idx, "undefined")

    def reset_finding_status(self):
        """
        Сбрасывает статус выбранной сработки в 'не обработано'.
        """
        idx = self._get_selected_index()
        if idx is None:
            return
        self._set_finding_status(idx, None)

    def _set_finding_status(self, idx: int, status: str | None):
        self.findings[idx]["status"] = status
        self._update_listbox_item_color(idx)
        # Сохраняем сразу, чтобы при любом выходе состояние было актуальным
        self._save_current_file_state()

    # --------- Скрытие/показ панели описания --------- #
    def toggle_description_panel(self):
        """
        Скрывает или показывает нижнюю панель с описанием сработки,
        при этом окно с кодом растягивается до низа окна.
        """
        if not hasattr(self, "bottom_frame") or self.bottom_frame is None:
            return

        if self.description_visible:
            # Скрываем панель описания
            try:
                self.bottom_frame.pack_forget()
            except tk.TclError:
                pass
            self.description_visible = False
            try:
                self.toggle_desc_btn.config(text="Показать описание")
            except tk.TclError:
                pass
        else:
            # Показываем панель описания обратно
            try:
                self.bottom_frame.pack(fill=tk.BOTH, expand=True, pady=(5, 0))
            except tk.TclError:
                pass
            self.description_visible = True
            try:
                self.toggle_desc_btn.config(text="Скрыть описание")
            except tk.TclError:
                pass

    # --------- Переключение темы (светлая / тёмная) --------- #
    def toggle_theme(self):
        """
        Переключает тему оформления между светлой и тёмной (VS Code‑like).
        """
        if self.is_dark_theme:
            self._apply_light_theme()
            self.is_dark_theme = False
            try:
                self.theme_btn.config(text="Тёмная тема (experimental)")
            except tk.TclError:
                pass
        else:
            self._apply_dark_theme()
            self.is_dark_theme = True
            try:
                self.theme_btn.config(text="Светлая тема")
            except tk.TclError:
                pass

    def _apply_light_theme(self):
        """
        Светлая тема (по умолчанию).
        """
        # Возвращаем исходную ttk тему (на Windows это обычно vista/xpnative)
        try:
            if self._default_ttk_theme:
                self._style.theme_use(self._default_ttk_theme)
        except tk.TclError:
            pass

        try:
            self.configure(bg="#f0f0f0")
        except tk.TclError:
            pass

        try:
            self._style.configure(".", background="#f0f0f0")
            self._style.configure("TFrame", background="#f0f0f0")
            self._style.configure("TLabel", background="#f0f0f0", foreground="black")
            # В светлой теме на нативных темах цвета кнопок могут игнорироваться — это ок.
            self._style.configure("TButton", foreground="black")

            self._style.configure(
                "Findings.Treeview",
                background="white",
                fieldbackground="white",
                foreground="black",
            )
            self._style.map(
                "Findings.Treeview",
                background=[("selected", "#cce5ff")],
                foreground=[("selected", "black")],
            )
        except tk.TclError:
            pass

        try:
            self.code_text.config(bg="white", fg="black", insertbackground="black")
        except tk.TclError:
            pass
        if getattr(self, "description_text", None) is not None:
            try:
                self.description_text.config(bg="white", fg="black", insertbackground="black")
            except tk.TclError:
                pass

        # Цвет подсветки фрагмента кода в светлой теме
        try:
            self.code_text.tag_configure("finding_highlight", background="#e6e6e6")
        except tk.TclError:
            pass

        # Обновляем цвета Rule id после применения темы
        for rule_id in self._rule_to_tree_item.keys():
            self._update_rule_item(rule_id)

    def _apply_dark_theme(self):
        """
        Тёмная тема, стилизованная под VS Code Dark.
        """
        bg = "#1e1e1e"
        fg = "#d4d4d4"
        panel_bg = "#252526"
        tree_bg = "#252526"
        tree_sel_bg = "#094771"
        btn_bg = "#3c3c3c"
        btn_active_bg = "#505050"

        # На Windows нативные ttk темы (vista/xpnative) не позволяют нормально перекрасить кнопки.
        # Поэтому для тёмной темы переключаемся на 'clam' (она полностью стилизуется).
        try:
            self._style.theme_use("clam")
        except tk.TclError:
            pass

        try:
            self.configure(bg=bg)
        except tk.TclError:
            pass

        try:
            # Базовые стили контейнеров
            self._style.configure(".", background=bg, foreground=fg)
            self._style.configure("TFrame", background=bg)
            self._style.configure("TLabel", background=bg, foreground=fg)

            # Кнопки (VS Code‑like)
            self._style.configure(
                "TButton",
                background=btn_bg,
                foreground=fg,
                bordercolor=btn_bg,
                focusthickness=1,
                focuscolor=tree_sel_bg,
                padding=(10, 5),
            )
            self._style.map(
                "TButton",
                background=[("active", btn_active_bg), ("pressed", btn_active_bg), ("disabled", btn_bg)],
                foreground=[("disabled", "#777777")],
            )

            # Ввод/текст (на будущее)
            self._style.configure("TEntry", fieldbackground=panel_bg, foreground=fg, background=panel_bg)
            self._style.configure("TCombobox", fieldbackground=panel_bg, foreground=fg, background=panel_bg)

            self._style.configure(
                "Findings.Treeview",
                background=tree_bg,
                fieldbackground=tree_bg,
                foreground=fg,
            )
            self._style.map(
                "Findings.Treeview",
                background=[("selected", tree_sel_bg)],
                foreground=[("selected", "white")],
            )

            self._style.configure("TPanedwindow", background=bg)
        except tk.TclError:
            pass

        try:
            self.code_text.config(bg=panel_bg, fg=fg, insertbackground=fg)
        except tk.TclError:
            pass
        if getattr(self, "description_text", None) is not None:
            try:
                self.description_text.config(bg=panel_bg, fg=fg, insertbackground=fg)
            except tk.TclError:
                pass

        # Цвет подсветки фрагмента кода в тёмной теме
        try:
            self.code_text.tag_configure("finding_highlight", background="#3a3a3a")
        except tk.TclError:
            pass

        # Обновляем цвета Rule id после применения темы
        for rule_id in self._rule_to_tree_item.keys():
            self._update_rule_item(rule_id)

    # --------- Формирование DOCX‑отчёта --------- #
    def _extract_cwe(self, finding: dict) -> str:
        """Извлекает CWE из ruleId или properties"""
        rule_id = str(finding.get("ruleId") or "")
        # Пытаемся найти CWE в ruleId (например, "CWE-79" или "CWE79")
        cwe_match = re.search(r'CWE[-\s]?(\d+)', rule_id, re.IGNORECASE)
        if cwe_match:
            return f"CWE-{cwe_match.group(1)}"
        return ""

    def _extract_recommendations_from_suggested_fix(self, suggested_fix: str) -> str:
        """Извлекает рекомендации из suggestedFix, останавливаясь на первом блоке кода"""
        if not suggested_fix:
            return ""
        
        # Преобразуем в строку, если еще не строка
        suggested_fix = str(suggested_fix).strip()
        if not suggested_fix:
            return ""
        
        lines = suggested_fix.split("\n")
        recommendations = []
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            
            # Проверяем начало блока кода - если встретили блок кода, останавливаемся
            if line_stripped.startswith("```"):
                # Останавливаемся на первом блоке кода - не включаем текст после него
                break
            
            # Добавляем строку, даже если она пустая (сохраняем структуру)
            # Но пропускаем полностью пустые строки в начале
            if line_stripped or (recommendations and recommendations[-1].strip()):
                recommendations.append(line_stripped if line_stripped else line)
        
        if recommendations:
            result = "\n".join(recommendations).strip()
            return result if result else ""
        return ""

    def _extract_recommendations(self, description: str) -> str:
        """Извлекает рекомендации из описания, останавливаясь на первом блоке кода"""
        if not description:
            return ""
        # Ищем раздел с рекомендациями
        lines = description.split("\n")
        recommendations = []
        in_recommendations = False
        
        i = 0
        while i < len(lines):
            line = lines[i]
            line_stripped = line.strip()
            line_lower = line_stripped.lower()
            
            # Проверяем начало блока кода - если встретили блок кода, останавливаемся
            if line_stripped.startswith("```"):
                # Останавливаемся на первом блоке кода - не включаем текст после него
                break
            
            # Ищем начало раздела рекомендаций
            if not in_recommendations:
                if "рекоменд" in line_lower or "suggested" in line_lower or "fix" in line_lower:
                    in_recommendations = True
                    # Пропускаем заголовок
                    if ":" in line:
                        i += 1
                        continue
            
            if in_recommendations:
                if line_stripped:
                    recommendations.append(line_stripped)
            
            i += 1
        
        if recommendations:
            return "\n".join(recommendations)
        return ""

    def _get_severity_order(self, severity: str) -> int:
        """Возвращает порядковый номер критичности для сортировки"""
        order_map = {
            "critical": 0,
            "high": 1,
            "middle": 2,
            "low": 3,
            "error": 4,
            "info": 5
        }
        return order_map.get(severity.lower() if severity else "", 99)

    def generate_doc_report(self):
        """
        Формирует DOCX‑отчёт по подтверждённым сработкам.
        """
        if not self._docx:
            messagebox.showerror(
                "Нет зависимости",
                "Для формирования DOC отчёта требуется пакет 'python-docx'.\n"
                "Установите его командой:\n\npip install python-docx",
            )
            return

        confirmed_findings = [f for f in self.findings if f.get("status") == "confirmed"]
        if not confirmed_findings:
            messagebox.showinfo("Нет данных", "Нет подтверждённых сработок для отчёта.")
            return

        # Сортируем по критичности (Critical, High, Middle, Low, Error, Info), затем по ruleId
        # Используем критичность из выпадающего списка для каждой сработки
        def _sort_key(f: dict):
            # Берем критичность из выпадающего списка (severity), если не установлена - используем Middle по умолчанию
            severity = f.get("severity")
            if not severity:
                # Если критичность не установлена, пытаемся определить из level
                level = f.get("level", "").lower()
                severity_mapping = {
                    "error": "Error",
                    "warning": "Middle",
                    "note": "Info"
                }
                severity = severity_mapping.get(level, "Middle")
            return (
                self._get_severity_order(severity),
                str(f.get("ruleId") or ""),
                str(f.get("uri") or ""),
                int(f.get("startLine") or 0),
            )

        # Применяем сортировку
        confirmed_findings.sort(key=_sort_key)

        # Выбор пути сохранения
        default_name = "appsec_report.docx"
        initial_dir = os.path.dirname(self.current_sarif_path) if self.current_sarif_path else ""
        save_path = filedialog.asksaveasfilename(
            title="Сохранить DOC отчёт",
            defaultextension=".docx",
            initialdir=initial_dir,
            initialfile=default_name,
            filetypes=[("Документы Word", "*.docx"), ("Все файлы", "*.*")],
        )
        if not save_path:
            return

        doc = self._docx.Document()  # type: ignore[call-arg]

        # Устанавливаем размер страницы A4
        try:
            from docx.shared import Inches, Mm  # type: ignore
            section = doc.sections[0]
            # A4 размер: 210 x 297 мм
            section.page_width = Mm(210)
            section.page_height = Mm(297)
        except Exception:
            # Если не удалось установить размер, продолжаем с настройками по умолчанию
            pass

        # Заголовок отчёта
        doc.add_heading("Отчёт по результатам анализа безопасности приложения", 0)

        # Предупреждение о том, что отчет не подлежит отправке заказчику
        warning_text = "ВНИМАНИЕ! ДАННЫЙ ОТЧЕТ НЕ ПОДЛЕЖИТ ОТПРАВКЕ ЗАКАЗЧИКУ. ОН ПРЕДНАЗНАЧЕН ТОЛЬКО ДЛЯ РЕДАКТИРОВАНИЯ И ТОЛЬКО ПОСЛЕ РЕДАКТИРОВАНИЯ МОЖЕТ БЫТЬ ОТПРАВЛЕН."
        warning_paragraph = doc.add_paragraph()
        warning_run = warning_paragraph.add_run(warning_text)
        warning_run.bold = True
        # Устанавливаем размер шрифта (12pt)
        try:
            from docx.shared import Pt  # type: ignore
            warning_run.font.size = Pt(12)
        except Exception:
            # Если не удалось установить размер, просто оставляем жирный текст
            pass
        # Добавляем пустую строку после предупреждения
        doc.add_paragraph()

        now = datetime.now().strftime("%d.%m.%Y %H:%M")
        app_name = os.path.basename(self.current_sarif_path) if self.current_sarif_path else "N/A"

        intro = (
            f"Настоящий отчёт подготовлен отделом прикладной безопасности (AppSec) по результатам "
            f"автоматизированного и экспертного анализа исходного кода.\n\n"
            f"Источник данных: SARIF‑отчёт «{app_name}».\n"
            f"Дата формирования отчёта: {now}."
        )
        for line in intro.split("\n"):
            doc.add_paragraph(line)

        doc.add_paragraph()

        # Краткая сводка
        total_confirmed = len(confirmed_findings)
        severity_counter: dict[str, int] = defaultdict(int)
        for f in confirmed_findings:
            severity = f.get("severity") or "Middle"
            severity_counter[severity] += 1

        doc.add_heading("1. Сводка по подтверждённым уязвимостям", level=1)
        doc.add_paragraph(f"Общее количество подтверждённых сработок: {total_confirmed}.")

        if severity_counter:
            p = doc.add_paragraph("Распределение по уровням критичности:")
            severity_order = ["Critical", "High", "Middle", "Low", "Error", "Info"]
            for sev in severity_order:
                if sev in severity_counter:
                    cnt = severity_counter[sev]
                    doc.add_paragraph(f"- {sev}: {cnt}", style="List Bullet")

        # Таблица с детализированными данными
        doc.add_heading("2. Таблица уязвимостей", level=1)

        # Создаём таблицу
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Grid Accent 1'
        
        # Устанавливаем ширину таблицы на всю страницу
        # Получаем ширину страницы (в twips, 1 дюйм = 1440 twips)
        section = doc.sections[0]
        # Ширина страницы минус поля (в twips)
        # Стандартная ширина страницы A4: 8.27 дюйма = 11906 twips
        # Стандартные поля: левое 1.25 дюйма, правое 1.25 дюйма
        # Рабочая ширина: 8.27 - 1.25 - 1.25 = 5.77 дюйма = 8309 twips
        available_width = int((section.page_width - section.left_margin - section.right_margin))
        
        # Устанавливаем ширину колонок (в twips)
        # Распределяем ширину пропорционально: № - 5%, Критичность - 8%, Тип - 12%, CWE - 10%, Файл - 15%, Описание - 25%, Рекомендации - 25%
        column_widths = [
            int(available_width * 0.05),      # №
            int(available_width * 0.08),      # Критичность
            int(available_width * 0.12),      # Тип уязвимости
            int(available_width * 0.10),      # CWE
            int(available_width * 0.15),      # Файл
            int(available_width * 0.25),      # Описание
            int(available_width * 0.25),      # Рекомендации
        ]
        
        for idx, width in enumerate(column_widths):
            try:
                table.columns[idx].width = width
            except Exception:
                # Если не удалось установить ширину, продолжаем
                pass

        # Заголовки таблицы
        header_cells = table.rows[0].cells
        header_cells[0].text = "№"
        header_cells[1].text = "Критичность"
        header_cells[2].text = "Тип уязвимости"
        header_cells[3].text = "CWE"
        header_cells[4].text = "Файл"
        header_cells[5].text = "Описание"
        header_cells[6].text = "Рекомендации"

        # Заполняем таблицу данными
        for idx, f in enumerate(confirmed_findings, start=1):
            row_cells = table.add_row().cells
            
            # №
            row_cells[0].text = str(idx)
            
            # Критичность (используем значение из выпадающего списка для каждой сработки)
            severity = f.get("severity") or "Middle"
            row_cells[1].text = severity
            
            # Тип уязвимости (ruleId)
            rule_id = f.get("ruleId") or "N/A"
            row_cells[2].text = rule_id
            
            # CWE
            cwe = self._extract_cwe(f)
            row_cells[3].text = cwe if cwe else "-"
            
            # Файл
            uri = f.get("uri") or "N/A"
            start_line = f.get("startLine")
            end_line = f.get("endLine")
            if start_line:
                file_info = f"{uri} (строки {start_line}"
                if end_line and end_line != start_line:
                    file_info += f"-{end_line}"
                file_info += ")"
            else:
                file_info = uri
            row_cells[4].text = file_info
            
            # Описание - используем только reasoning
            reasoning = f.get("reasoning") or ""
            if reasoning:
                row_cells[5].text = reasoning.strip()
            else:
                row_cells[5].text = "-"
            
            # Рекомендации - используем только suggestedFix, обрезаем до первого блока кода
            suggested_fix = f.get("suggestedFix")
            
            # Обрабатываем разные форматы suggestedFix
            if isinstance(suggested_fix, dict):
                suggested_fix = suggested_fix.get("text", "") or suggested_fix.get("description", "") or ""
            elif isinstance(suggested_fix, list):
                if len(suggested_fix) > 0:
                    if isinstance(suggested_fix[0], dict):
                        texts = []
                        for item in suggested_fix:
                            text = item.get("text") or item.get("description") or str(item)
                            if text:
                                texts.append(str(text))
                        suggested_fix = "\n".join(texts)
                    else:
                        suggested_fix = "\n".join(str(item) for item in suggested_fix if item)
                else:
                    suggested_fix = ""
            
            # Преобразуем в строку
            if not isinstance(suggested_fix, str):
                suggested_fix = str(suggested_fix) if suggested_fix else ""
            
            if suggested_fix and suggested_fix.strip():
                # Обрезаем до первого блока кода
                recommendations = self._extract_recommendations_from_suggested_fix(suggested_fix.strip())
                if recommendations and recommendations.strip():
                    row_cells[6].text = recommendations
                else:
                    # Если после обработки ничего не осталось, выводим исходный текст до блока кода
                    lines = suggested_fix.split("\n")
                    text_before_code = []
                    for line in lines:
                        if line.strip().startswith("```"):
                            break
                        if line.strip():  # Пропускаем пустые строки в начале
                            text_before_code.append(line.strip())
                    
                    result_text = "\n".join(text_before_code).strip()
                    row_cells[6].text = result_text if result_text else "-"
            else:
                row_cells[6].text = "-"

        try:
            doc.save(save_path)
        except Exception as e:
            messagebox.showerror("Ошибка сохранения", f"Не удалось сохранить DOC отчёт:\n{e}")
            return

        messagebox.showinfo("Готово", f"DOC отчёт успешно сохранён:\n{save_path}")


class ComparisonWindow(tk.Toplevel):
    """
    Окно для сравнения двух JSON файлов состояния одного SARIF файла.
    Показывает только обработанные сработки (те, что отличаются между двумя JSON).
    """

    def __init__(self, parent, findings, state1, state2, filename1, filename2, sarif_path):
        super().__init__(parent)
        self.title(f"Сравнение JSON файлов: {filename1} vs {filename2}")
        self.geometry("1100x700")

        self.findings = findings
        self.state1 = state1
        self.state2 = state2
        self.filename1 = filename1
        self.filename2 = filename2
        self.sarif_path = sarif_path

        # Находим различия
        self.diff_findings = self._find_differences()

        # Опциональная поддержка Markdown в описании
        self._md = None
        self._HtmlFrame = None
        try:
            import markdown as _markdown
            self._md = _markdown
        except Exception:
            self._md = None
        try:
            from tkinterweb import HtmlFrame as _HtmlFrame
            self._HtmlFrame = _HtmlFrame
        except Exception:
            self._HtmlFrame = None

        self._build_ui()
        self._populate_tree()

    def _make_finding_key(self, f: dict) -> str:
        """Делает стабильный ключ для идентификации сработки."""
        return "|".join(
            [
                str(f.get("ruleId") or ""),
                str(f.get("uri") or ""),
                str(f.get("startLine") or ""),
                str(f.get("endLine") or ""),
                str(f.get("message") or ""),
            ]
        )

    def _find_differences(self):
        """
        Находит сработки, которые отличаются между двумя состояниями.
        Возвращает список словарей с информацией о различиях.
        Показывает только сработки с различиями в статусе или критичности.
        Исключает сработки, которые не обработаны в обоих файлах.
        """
        diff_findings = []
        
        # Создаем множество всех ключей из обоих состояний
        all_keys = set(self.state1.keys()) | set(self.state2.keys())
        
        for key in all_keys:
            entry1 = self.state1.get(key)
            entry2 = self.state2.get(key)
            
            # Определяем статусы и критичность из обоих состояний
            status1 = None
            severity1 = None
            if isinstance(entry1, dict):
                status1 = entry1.get("status")
                severity1 = entry1.get("severity")
            elif isinstance(entry1, str):
                status1 = entry1
            
            status2 = None
            severity2 = None
            if isinstance(entry2, dict):
                status2 = entry2.get("status")
                severity2 = entry2.get("severity")
            elif isinstance(entry2, str):
                status2 = entry2
            
            # Пропускаем сработки, которые не обработаны в обоих файлах
            if status1 is None and status2 is None:
                continue
            
            # Проверяем, есть ли различия в статусе или критичности
            status_differs = status1 != status2
            severity_differs = severity1 != severity2
            
            # Если есть различия, добавляем сработку
            if status_differs or severity_differs:
                # Находим соответствующую сработку
                for finding in self.findings:
                    finding_key = self._make_finding_key(finding)
                    if finding_key == key:
                        # Создаем копию сработки с информацией о различиях
                        diff_finding = finding.copy()
                        
                        diff_finding["status1"] = status1
                        diff_finding["severity1"] = severity1
                        diff_finding["status2"] = status2
                        diff_finding["severity2"] = severity2
                        
                        diff_findings.append(diff_finding)
                        break
        
        return diff_findings

    def _build_ui(self):
        """Создает интерфейс окна сравнения."""
        # Верхняя панель с информацией
        top_frame = ttk.Frame(self)
        top_frame.pack(side=tk.TOP, fill=tk.X, padx=5, pady=5)

        info_label = ttk.Label(
            top_frame,
            text=f"Показаны только обработанные сработки, отличающиеся между файлами.\n"
                 f"Файл 1: {self.filename1} | Файл 2: {self.filename2}"
        )
        info_label.pack(side=tk.LEFT)

        # Основная панель с разделением
        main_pane = ttk.Panedwindow(self, orient=tk.HORIZONTAL)
        main_pane.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Левая панель – список сработок
        left_frame = ttk.Frame(main_pane)
        main_pane.add(left_frame, weight=1)

        left_label = ttk.Label(left_frame, text="Различия")
        left_label.pack(side=tk.TOP, anchor="w", pady=(0, 5))

        self.findings_tree = ttk.Treeview(
            left_frame,
            show="tree",
            selectmode="browse",
        )
        self.findings_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.findings_tree.bind("<<TreeviewSelect>>", self.on_finding_selected)

        scrollbar_left = ttk.Scrollbar(left_frame, orient=tk.VERTICAL, command=self.findings_tree.yview)
        scrollbar_left.pack(side=tk.RIGHT, fill=tk.Y)
        self.findings_tree.config(yscrollcommand=scrollbar_left.set)

        self._tree_item_to_finding_index = {}

        # Правая панель – код и описание
        right_frame = ttk.Frame(main_pane)
        main_pane.add(right_frame, weight=3)

        # Верхняя часть справа – код
        self.code_label = ttk.Label(right_frame, text="Фрагмент кода")
        self.code_label.pack(side=tk.TOP, anchor="w")

        self.code_text = ScrolledText(right_frame, height=18, wrap=tk.NONE)
        self.code_text.pack(fill=tk.BOTH, expand=True)
        self.code_text.config(state=tk.DISABLED)
        try:
            self.code_text.tag_configure("finding_highlight", background="#e6e6e6")
        except tk.TclError:
            pass

        # Нижняя часть – описание и информация о различиях
        bottom_frame = ttk.Frame(right_frame)
        bottom_frame.pack(fill=tk.BOTH, expand=True, pady=(5, 0))

        # Информация о различиях
        diff_label = ttk.Label(bottom_frame, text="Различия в обработке")
        diff_label.pack(anchor="w", pady=(5, 0))

        self.diff_text = ScrolledText(bottom_frame, height=8, wrap=tk.WORD)
        self.diff_text.pack(fill=tk.BOTH, expand=True)
        self.diff_text.config(state=tk.DISABLED)

        # Описание
        desc_label = ttk.Label(bottom_frame, text="Описание уязвимости")
        desc_label.pack(anchor="w", pady=(5, 0))

        self.description_widget_kind = "text"
        self.description_text = None
        self.description_html = None
        if self._md and self._HtmlFrame:
            self.description_widget_kind = "html"
            self.description_html = self._HtmlFrame(bottom_frame, horizontal_scrollbar="auto")
            self.description_html.pack(fill=tk.BOTH, expand=True)
        else:
            self.description_text = ScrolledText(bottom_frame, height=10, wrap=tk.WORD)
            self.description_text.pack(fill=tk.BOTH, expand=True)
            self.description_text.config(state=tk.DISABLED)

    def _populate_tree(self):
        """Заполняет дерево различиями."""
        for item in self.findings_tree.get_children(""):
            self.findings_tree.delete(item)
        self._tree_item_to_finding_index.clear()

        # Группируем по ruleId
        grouped: dict[str, list[int]] = defaultdict(list)
        for idx, f in enumerate(self.diff_findings):
            grouped[f.get("ruleId") or "N/A"].append(idx)

        for rule_id in sorted(grouped.keys(), key=lambda s: str(s).lower()):
            count = len(grouped[rule_id])
            rule_item = self.findings_tree.insert("", tk.END, text=f"{rule_id} ({count})", open=True)
            for idx in grouped[rule_id]:
                f = self.diff_findings[idx]
                title = f"{idx + 1}. [{f['level'].upper()}] {f['message']}"
                child = self.findings_tree.insert(rule_item, tk.END, text=title, open=False)
                self._tree_item_to_finding_index[child] = idx
                self.diff_findings[idx]["tree_item"] = child

                # Цветовая индикация различий (статус или критичность)
                status1 = f.get("status1")
                status2 = f.get("status2")
                severity1 = f.get("severity1")
                severity2 = f.get("severity2")
                if status1 != status2 or severity1 != severity2:
                    tag = "status_diff"
                    try:
                        self.findings_tree.tag_configure(tag, background="#ffe6cc")  # светло-оранжевый
                        self.findings_tree.item(child, tags=(tag,))
                    except tk.TclError:
                        pass

    def on_finding_selected(self, event=None):
        """Обработчик выбора сработки."""
        selection = self.findings_tree.selection()
        if not selection:
            return
        item = selection[0]

        if item not in self._tree_item_to_finding_index:
            return

        idx = self._tree_item_to_finding_index[item]
        finding = self.diff_findings[idx]

        # Показываем код
        self._show_code_for_finding(finding)

        # Показываем описание
        self._set_description(finding.get("description") or "Описание отсутствует.")

        # Показываем различия
        self._show_differences(finding)

    def _show_code_for_finding(self, finding: dict):
        """Показывает код для сработки."""
        if finding.get("uri"):
            try:
                self.code_label.config(text=f"Фрагмент кода — {finding['uri']}")
            except tk.TclError:
                pass
        else:
            try:
                self.code_label.config(text="Фрагмент кода")
            except tk.TclError:
                pass

        header_parts = []
        if finding["uri"]:
            header_parts.append(f"Файл: {finding['uri']}")
        if finding["startLine"]:
            header_parts.append(f"Строки: {finding['startLine']}–{finding['endLine']}")

        header = " | ".join(header_parts)
        if finding.get("full_code"):
            self._set_code_with_highlight(
                widget=self.code_text,
                header=header,
                full_code=finding["full_code"],
                start_line=finding.get("startLine"),
                end_line=finding.get("endLine"),
                uri=finding.get("uri"),
            )
            return

        code_to_show = finding.get("snippet") or "Фрагмент кода недоступен."
        if header:
            code_to_show = header + "\n\n" + code_to_show
        self._set_text(self.code_text, code_to_show)

    def _set_text(self, widget: ScrolledText, text: str):
        """Устанавливает текст в виджет."""
        if widget is None:
            return
        widget.config(state=tk.NORMAL)
        widget.delete("1.0", tk.END)
        widget.insert(tk.END, text)
        widget.config(state=tk.DISABLED)

    def _set_code_with_highlight(
        self,
        widget: ScrolledText,
        header: str,
        full_code: str,
        start_line: int | None,
        end_line: int | None,
        uri: str | None,
    ):
        """Показывает код с подсветкой (упрощенная версия из основного класса)."""
        lines = full_code.splitlines()
        max_no = max(1, len(lines))
        width = len(str(max_no))

        hl_start = None
        hl_end = None
        if isinstance(start_line, int) and start_line >= 1:
            hl_start = start_line
        if isinstance(end_line, int) and end_line >= 1:
            hl_end = end_line
        if hl_start is not None and hl_end is None:
            hl_end = hl_start
        if hl_start is not None and hl_end is not None and hl_end < hl_start:
            hl_start, hl_end = hl_end, hl_start

        widget.config(state=tk.NORMAL)
        widget.delete("1.0", tk.END)
        try:
            widget.tag_remove("finding_highlight", "1.0", tk.END)
        except tk.TclError:
            pass

        text_parts = []
        header_lines_count = 0
        if header:
            text_parts.append(header)
            text_parts.append("")
            header_lines_count = 2

        numbered_lines = [f"{i:{width}d}: {line}" for i, line in enumerate(lines, start=1)]
        text_parts.extend(numbered_lines)

        full_text = "\n".join(text_parts) + ("\n" if lines else "")
        widget.insert("1.0", full_text)

        if hl_start is not None and hl_end is not None:
            code_start_widget_line = header_lines_count + 1
            for i in range(hl_start, hl_end + 1):
                widget_line = code_start_widget_line + (i - 1)
                start_index = f"{widget_line}.0"
                end_index = f"{widget_line}.end"
                try:
                    widget.tag_add("finding_highlight", start_index, end_index)
                except tk.TclError:
                    pass

            first_widget_line = code_start_widget_line + (hl_start - 1)
            try:
                widget.see(f"{first_widget_line}.0")
            except tk.TclError:
                pass

        widget.config(state=tk.DISABLED)

    def _show_differences(self, finding: dict):
        """Показывает различия между двумя состояниями."""
        status1 = finding.get("status1")
        severity1 = finding.get("severity1")
        status2 = finding.get("status2")
        severity2 = finding.get("severity2")

        diff_lines = []
        diff_lines.append(f"Файл 1 ({self.filename1}):")
        diff_lines.append(f"  Статус: {status1 or 'не обработано'}")
        diff_lines.append(f"  Критичность: {severity1 or 'не установлена'}")
        diff_lines.append("")
        diff_lines.append(f"Файл 2 ({self.filename2}):")
        diff_lines.append(f"  Статус: {status2 or 'не обработано'}")
        diff_lines.append(f"  Критичность: {severity2 or 'не установлена'}")
        diff_lines.append("")
        
        # Показываем, что именно отличается
        differences = []
        if status1 != status2:
            differences.append("Статус")
        if severity1 != severity2:
            differences.append("Критичность")
        
        if differences:
            diff_lines.append(f"Отличается: {', '.join(differences)}")

        diff_text = "\n".join(diff_lines)
        self._set_text(self.diff_text, diff_text)

    def _set_description(self, markdown_text: str):
        """Устанавливает описание (упрощенная версия из основного класса)."""
        if self.description_widget_kind == "html" and self.description_html is not None:
            try:
                html = self._md.markdown(
                    markdown_text,
                    extensions=["fenced_code", "tables", "toc"],
                    output_format="html5",
                )
                css = (
                    "body{font-family:Segoe UI,Arial,sans-serif;font-size:10pt;}"
                    "pre,code{font-family:Consolas,Menlo,monospace;font-size:9pt;}"
                    "pre{background:#f6f6f6;padding:8px;border-radius:6px;}"
                    "code{background:#f0f0f0;padding:1px 3px;border-radius:4px;}"
                    "table{border-collapse:collapse;}td,th{border:1px solid #ddd;padding:6px;}"
                )
                html_doc = (
                    "<html><head><meta charset='utf-8'>"
                    "<style>" + css + "</style></head><body>"
                    + html + "</body></html>"
                )
                self.description_html.load_html(html_doc)
            except Exception:
                if self.description_text is not None:
                    self._set_text(self.description_text, markdown_text)
            return

        if self.description_text is not None:
            self._set_text(self.description_text, markdown_text)


def main():
    app = SarifViewer()
    app.mainloop()


if __name__ == "__main__":
    main()


